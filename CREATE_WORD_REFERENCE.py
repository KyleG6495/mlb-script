#!/usr/bin/env python3
"""
Creates a Word document version of the Elite DFS Workflow Reference Guide
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def create_word_reference_guide():
    """Create a comprehensive Word document reference guide"""
    
    try:
        # Create document
        doc = Document()
        
        # Set up styles
        styles = doc.styles
        
        # Title
        title = doc.add_heading('ELITE DFS WORKFLOW REFERENCE GUIDE', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = doc.add_paragraph('Complete Script Reference & File Location Guide')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_format = subtitle.runs[0].font
        subtitle_format.size = Pt(14)
        subtitle_format.italic = True
        
        # Date
        date_para = doc.add_paragraph('Created: August 22, 2025')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # Table of Contents
        doc.add_heading('TABLE OF CONTENTS', level=1)
        toc_items = [
            '1. COMPLETE_DAILY_WORKFLOW.bat - Main Daily Process',
            '2. ENHANCED_DAILY_WORKFLOW.bat - Optional Enhancements',
            '3. File Locations and Output Directory Structure', 
            '4. Quick Reference Commands',
            '5. System Requirements and Notes'
        ]
        
        for item in toc_items:
            doc.add_paragraph(item, style='List Number')
        
        doc.add_page_break()
        
        # Section 1: Complete Daily Workflow
        doc.add_heading('1. COMPLETE_DAILY_WORKFLOW.bat - MAIN DAILY PROCESS', level=1)
        
        doc.add_paragraph('File Location: C:\\Users\\kgone\\OneDrive\\Personal_Information\\MLB\\Scripts\\COMPLETE_DAILY_WORKFLOW.bat')
        
        workflow_steps = [
            {
                'step': 'STEP 1: CORE DATA PIPELINE',
                'script': '1_DATA_PIPELINE.bat',
                'location': 'DAILY_RUNNERS\\',
                'purpose': 'Downloads and processes base MLB data',
                'outputs': [
                    '../data/hitter_games_today.csv',
                    '../data/pitcher_games_today.csv', 
                    '../data/team_stats_today.csv',
                    '../data/weather_today.csv'
                ]
            },
            {
                'step': 'STEP 2: BASE DFS MODELS',
                'script': '2_DFS_MODELS.bat',
                'location': 'DAILY_RUNNERS\\',
                'purpose': 'Builds initial projections and models',
                'outputs': [
                    '../data/hitter_projections_today.csv',
                    '../data/pitcher_projections_today.csv',
                    '../data/base_lineups_today.csv'
                ]
            },
            {
                'step': 'STEP 3: VEGAS INTEGRATION',
                'script': 'VEGAS_ODDS_INTEGRATOR.py',
                'location': 'Scripts\\',
                'purpose': 'Integrates Vegas odds and game totals',
                'outputs': [
                    'data/vegas_odds_today.csv',
                    'data/game_totals_today.csv',
                    'data/enhanced_projections_with_vegas.csv'
                ]
            },
            {
                'step': 'STEP 4: OWNERSHIP PROJECTIONS',
                'script': 'ADVANCED_OWNERSHIP_PROJECTIONS.py',
                'location': 'Scripts\\',
                'purpose': 'Generates sophisticated ownership modeling',
                'outputs': [
                    'data/advanced_ownership_projections_[timestamp].csv',
                    'data/ownership_leverage_analysis.csv',
                    'data/chalk_fade_opportunities.csv'
                ]
            },
            {
                'step': 'STEP 5: ELITE LINEUP GENERATION',
                'script': 'ELITE_TOURNAMENT_WITH_OWNERSHIP.py',
                'location': 'Scripts\\',
                'purpose': 'Creates tournament-focused lineups with ownership consideration',
                'outputs': [
                    'fd_current_slate/elite_tournament_lineups_[timestamp].csv',
                    'data/tournament_strategy_analysis.csv',
                    'data/lineup_construction_log.csv'
                ]
            },
            {
                'step': 'STEP 6: ENHANCED MODELS & OPTIMIZATION',
                'script': '4_ENHANCED_MODELS.bat + Additional Scripts',
                'location': 'DAILY_RUNNERS\\ + Scripts\\',
                'purpose': 'Advanced modeling and optimization',
                'outputs': [
                    'data/enhanced_projections_[timestamp].csv',
                    'data/correlation_matrix_today.csv',
                    'data/team_stack_analysis_[timestamp].csv',
                    'data/correlation_analysis_[timestamp].csv',
                    'data/leverage_analysis_[timestamp].csv'
                ]
            },
            {
                'step': 'STEP 7: REAL-TIME DATA INTEGRATION',
                'script': 'fetch_weather_data.py, fetch_rotowire_lineups_enhanced.py, fetch_live_scores.py',
                'location': 'Scripts\\',
                'purpose': 'Fetches live weather, lineups, and scores',
                'outputs': [
                    'data/weather_updates_today.json',
                    'temp_lineup_data.json',
                    'live_scores_today.json',
                    'confirmed_lineups_today.csv'
                ]
            },
            {
                'step': 'STEP 8: UMPIRE & EDGE ANALYSIS',
                'script': 'umpire_impact_analyzer.py, PARK_FACTOR_INTEGRATION.py',
                'location': 'Scripts\\',
                'purpose': 'Analyzes umpire impacts and park factors',
                'outputs': [
                    'data/umpire_analysis_today.csv',
                    'data/park_factors_today.csv'
                ]
            },
            {
                'step': 'STEP 9: FINAL LINEUP OPTIMIZATION',
                'script': 'FINAL_LINEUP_OPTIMIZER.py, ELITE_LINEUP_SELECTOR.py',
                'location': 'Scripts\\',
                'purpose': 'Final lineup optimization and selection',
                'outputs': [
                    'data/optimized_lineups_[timestamp].csv',
                    'fd_current_slate/selected_lineups_[timestamp].csv'
                ]
            },
            {
                'step': 'STEP 10: ADVANCED DFS INTEGRATION',
                'script': 'ADVANCED_DFS_INTEGRATOR.py',
                'location': 'Scripts\\',
                'purpose': 'System-wide DFS integration',
                'outputs': [
                    'fd_current_slate/integrated_lineups_[timestamp].csv',
                    'data/integration_summary_[timestamp].csv'
                ]
            },
            {
                'step': 'STEP 11: CONTEST-SPECIFIC EXPORT',
                'script': 'EXPORT_SELECTED_LINEUPS.py',
                'location': 'Scripts\\',
                'purpose': 'Creates FanDuel-ready lineup files',
                'outputs': [
                    'fd_current_slate/RECOMMENDED_cash_games_[timestamp].csv',
                    'fd_current_slate/RECOMMENDED_small_tournaments_[timestamp].csv',
                    'fd_current_slate/RECOMMENDED_large_tournaments_[timestamp].csv'
                ]
            },
            {
                'step': 'STEP 12: LAUNCH ELITE DASHBOARD',
                'script': 'COMPLETE_ELITE_DFS_DASHBOARD.py',
                'location': 'Scripts\\',
                'purpose': 'Interactive 7-tab dashboard for analysis and management',
                'outputs': [
                    'Dashboard interface with file management and export capabilities'
                ]
            }
        ]
        
        for step_info in workflow_steps:
            doc.add_heading(step_info['step'], level=2)
            
            # Create table for step details
            table = doc.add_table(rows=4, cols=2)
            table.style = 'Table Grid'
            
            # Headers and data
            table.cell(0, 0).text = 'Script(s):'
            table.cell(0, 1).text = step_info['script']
            table.cell(1, 0).text = 'Location:'
            table.cell(1, 1).text = step_info['location']
            table.cell(2, 0).text = 'Purpose:'
            table.cell(2, 1).text = step_info['purpose']
            table.cell(3, 0).text = 'Output Files:'
            table.cell(3, 1).text = '\\n'.join(step_info['outputs'])
            
            doc.add_paragraph()  # Space after table
        
        doc.add_page_break()
        
        # Section 2: Enhanced Workflow
        doc.add_heading('2. ENHANCED_DAILY_WORKFLOW.bat - OPTIONAL ENHANCEMENTS', level=1)
        
        doc.add_paragraph('File Location: C:\\Users\\kgone\\OneDrive\\Personal_Information\\MLB\\Scripts\\ENHANCED_DAILY_WORKFLOW.bat')
        
        enhanced_steps = [
            {
                'step': 'ENHANCED 1: VALIDATION & ANALYSIS',
                'scripts': [
                    'BACKTEST_VALIDATOR.py',
                    'QUALITY_VS_QUANTITY_ANALYSIS.py', 
                    'LINEUP_WORKFLOW_SUMMARY.py'
                ],
                'outputs': [
                    'data/backtest_results_[timestamp].csv',
                    'data/quality_analysis_[timestamp].csv',
                    'Console summary of lineup process'
                ]
            },
            {
                'step': 'ENHANCED 2: COMPETITIVE INTELLIGENCE',
                'scripts': [
                    'SABERSIM_COMPARISON.py',
                    'DFS_INDUSTRY_TRUTH.py'
                ],
                'outputs': [
                    'Console analysis of competitive positioning',
                    'Console analysis of industry practices'
                ]
            },
            {
                'step': 'ENHANCED 3: EXPANSION PLANNING',
                'scripts': [
                    'LINEUP_EXPANSION_ENGINE.py'
                ],
                'outputs': [
                    'data/expansion_analysis_[timestamp].csv',
                    'Console recommendations for lineup scaling'
                ]
            },
            {
                'step': 'ENHANCED 4: FINAL SYSTEM CHECK',
                'scripts': [
                    'SYSTEM_STATUS_CHECK.py'
                ],
                'outputs': [
                    'data/system_status_[timestamp].csv',
                    'Console system health report'
                ]
            }
        ]
        
        for enhanced_info in enhanced_steps:
            doc.add_heading(enhanced_info['step'], level=2)
            
            doc.add_paragraph('Scripts:', style='List Bullet')
            for script in enhanced_info['scripts']:
                doc.add_paragraph(f'• {script}', style='List Bullet 2')
            
            doc.add_paragraph('Outputs:', style='List Bullet')
            for output in enhanced_info['outputs']:
                doc.add_paragraph(f'• {output}', style='List Bullet 2')
        
        doc.add_page_break()
        
        # Section 3: File Locations
        doc.add_heading('3. FILE LOCATIONS AND DIRECTORY STRUCTURE', level=1)
        
        doc.add_heading('Main Directories:', level=2)
        directories = [
            'Scripts: C:\\Users\\kgone\\OneDrive\\Personal_Information\\MLB\\Scripts\\',
            'Data: C:\\Users\\kgone\\OneDrive\\Personal_Information\\MLB\\data\\',
            'FanDuel Slate: C:\\Users\\kgone\\OneDrive\\Personal_Information\\MLB\\fd_current_slate\\',
            'Daily Runners: C:\\Users\\kgone\\OneDrive\\Personal_Information\\MLB\\Scripts\\DAILY_RUNNERS\\'
        ]
        
        for directory in directories:
            doc.add_paragraph(directory, style='List Bullet')
        
        doc.add_heading('Key Output Files by Category:', level=2)
        
        file_categories = [
            {
                'category': 'PROJECTIONS & MODELS:',
                'files': [
                    'data/hitter_projections_today.csv',
                    'data/pitcher_projections_today.csv',
                    'data/enhanced_projections_[timestamp].csv',
                    'data/vegas_odds_today.csv'
                ]
            },
            {
                'category': 'OWNERSHIP & STRATEGY:',
                'files': [
                    'data/advanced_ownership_projections_[timestamp].csv',
                    'data/ownership_leverage_analysis.csv',
                    'data/chalk_fade_opportunities.csv'
                ]
            },
            {
                'category': 'LINEUPS & EXPORTS:',
                'files': [
                    'fd_current_slate/elite_tournament_lineups_[timestamp].csv',
                    'fd_current_slate/RECOMMENDED_cash_games_[timestamp].csv',
                    'fd_current_slate/RECOMMENDED_small_tournaments_[timestamp].csv',
                    'fd_current_slate/RECOMMENDED_large_tournaments_[timestamp].csv'
                ]
            },
            {
                'category': 'ANALYSIS & VALIDATION:',
                'files': [
                    'data/correlation_analysis_[timestamp].csv',
                    'data/team_stack_analysis_[timestamp].csv',
                    'data/umpire_analysis_today.csv',
                    'data/backtest_results_[timestamp].csv'
                ]
            },
            {
                'category': 'REAL-TIME DATA:',
                'files': [
                    'temp_lineup_data.json',
                    'live_scores_today.json',
                    'data/weather_updates_today.json',
                    'confirmed_lineups_today.csv'
                ]
            }
        ]
        
        for category_info in file_categories:
            doc.add_heading(category_info['category'], level=3)
            for file_path in category_info['files']:
                doc.add_paragraph(f'• {file_path}', style='List Bullet')
        
        doc.add_page_break()
        
        # Section 4: Quick Reference
        doc.add_heading('4. QUICK REFERENCE COMMANDS', level=1)
        
        commands = [
            {
                'title': 'To Run Complete Daily Process:',
                'commands': [
                    'cd "C:\\Users\\kgone\\OneDrive\\Personal_Information\\MLB\\Scripts"',
                    '.\\COMPLETE_DAILY_WORKFLOW.bat'
                ]
            },
            {
                'title': 'To Run Enhanced Validation:',
                'commands': [
                    'cd "C:\\Users\\kgone\\OneDrive\\Personal_Information\\MLB\\Scripts"',
                    '.\\ENHANCED_DAILY_WORKFLOW.bat'
                ]
            },
            {
                'title': 'To Launch Dashboard Only:',
                'commands': [
                    'cd "C:\\Users\\kgone\\OneDrive\\Personal_Information\\MLB\\Scripts"',
                    'python COMPLETE_ELITE_DFS_DASHBOARD.py'
                ]
            },
            {
                'title': 'To Export Lineups Only:',
                'commands': [
                    'cd "C:\\Users\\kgone\\OneDrive\\Personal_Information\\MLB\\Scripts"',
                    'python EXPORT_SELECTED_LINEUPS.py'
                ]
            },
            {
                'title': 'To Check Latest Files:',
                'commands': [
                    'cd "C:\\Users\\kgone\\OneDrive\\Personal_Information\\MLB\\fd_current_slate"',
                    'dir *.csv /O:D'
                ]
            }
        ]
        
        for command_info in commands:
            doc.add_heading(command_info['title'], level=2)
            for cmd in command_info['commands']:
                para = doc.add_paragraph(cmd)
                para.style = 'Intense Quote'
        
        # Section 5: System Requirements
        doc.add_heading('5. SYSTEM REQUIREMENTS AND NOTES', level=1)
        
        doc.add_heading('Important Notes:', level=2)
        notes = [
            'All timestamps use format: YYYYMMDD_HHMMSS',
            'FanDuel CSV files are ready for direct upload',
            'Dashboard provides file management and preview capabilities',
            'Enhanced workflow is optional but recommended for validation',
            'Scripts automatically create output directories if needed'
        ]
        
        for note in notes:
            doc.add_paragraph(f'• {note}', style='List Bullet')
        
        doc.add_heading('System Requirements:', level=2)
        requirements = [
            'Python 3.8+ with required packages',
            'Internet connection for real-time data',
            'Minimum 2GB free disk space for daily files',
            'Windows PowerShell for batch file execution'
        ]
        
        for req in requirements:
            doc.add_paragraph(f'• {req}', style='List Bullet')
        
        doc.add_heading('Support Files Created:', level=2)
        support_files = [
            'WORKFLOW_SUMMARY.py - Displays process overview',
            'HOW_TO_GET_LINEUPS.py - Step-by-step lineup access guide',
            'SABERSIM_COMPARISON.py - Competitive analysis',
            'DFS_INDUSTRY_TRUTH.py - Industry insight analysis'
        ]
        
        for support_file in support_files:
            doc.add_paragraph(f'• {support_file}', style='List Bullet')
        
        # Save document
        doc_path = 'ELITE_DFS_WORKFLOW_REFERENCE.docx'
        doc.save(doc_path)
        
        return doc_path
        
    except Exception as e:
        print(f"Error creating Word document: {e}")
        print("Creating text version instead...")
        return None

def main():
    print("🔄 Creating Elite DFS Workflow Reference Guide...")
    
    try:
        doc_path = create_word_reference_guide()
        if doc_path:
            print(f"✅ Word document created: {doc_path}")
            print("📁 You can copy this file to your desktop")
        else:
            print("❌ Word document creation failed")
            print("✅ Text version available: ELITE_DFS_WORKFLOW_REFERENCE.txt")
    
    except ImportError:
        print("📝 python-docx not installed. Creating text version only.")
        print("✅ Text version available: ELITE_DFS_WORKFLOW_REFERENCE.txt")
        print("💡 To create Word doc, install: pip install python-docx")
    
    print()
    print("📋 REFERENCE GUIDE CONTENTS:")
    print("  • Complete workflow breakdown (12 steps)")
    print("  • Enhanced workflow details (4 additional steps)")  
    print("  • All script locations and file outputs")
    print("  • Quick reference commands")
    print("  • System requirements and notes")
    print()
    print("🎯 You can now reference this guide for any specific questions!")

if __name__ == "__main__":
    main()
